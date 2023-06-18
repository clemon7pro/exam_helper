#!/usr/bin/python
# -*- coding: UTF-8 -*-
from typing import List
from docx import Document
import openpyxl
import math

from .question import *

class DocWriter(object):
    
    def __init__(self, file: str, questions: List[Question]) -> None:
        self.__doc = Document()
        self.__file = file
        self.__questions = questions
        
    def __write_single(self, id: int, question: Question) -> None:
        self.__doc.add_paragraph("第" + str(id) + "题. " + question.question_stem)
        answer = "#err"
        
        match question.type:
            case QuestionType.single_choice | QuestionType.multiple_choice:
                col_len = 2
                row_len = int(math.ceil(len(question.options) / col_len))
                table = self.__doc.add_table(rows=row_len, cols=col_len)
                
                for i, v in enumerate(question.options):
                    table.cell(int(divmod(i, col_len)[0]), int(divmod(i, col_len)[1])).text = chr(65+i)+". " + v
                    
                answer = "".join(question.answer)
            case QuestionType.binary_choice:
                answer = "对" if question.answer else "错"
            case QuestionType.short_answer:
                answer = question.answer
                
        self.__doc.add_paragraph("答案: " + answer)
    
    def write(self) -> None:
        single_choice_questions = list(filter(lambda item: item.type == QuestionType.single_choice, self.__questions))
        multiple_choice_questions = list(filter(lambda item: item.type == QuestionType.multiple_choice, self.__questions))
        binary_choice_questions = list(filter(lambda item: item.type == QuestionType.binary_choice, self.__questions))
        short_answer_questions = list(filter(lambda item: item.type == QuestionType.short_answer, self.__questions))
        
        id = 1
        if len(single_choice_questions):
            self.__doc.add_paragraph("单选题: ")
            for question in single_choice_questions:
                self.__write_single(id, question)
                self.__doc.add_paragraph()
                id += 1
                
        if len(multiple_choice_questions):
            self.__doc.add_paragraph("多选题: ")
            for question in multiple_choice_questions:
                self.__write_single(id, question)
                self.__doc.add_paragraph()
                id += 1
                
        if len(binary_choice_questions):
            self.__doc.add_paragraph("判断题: ")
            for question in binary_choice_questions:
                self.__write_single(id, question)
                self.__doc.add_paragraph()
                id += 1
        
        if len(short_answer_questions):
            self.__doc.add_paragraph("简答题: ")
            for question in short_answer_questions:
                self.__write_single(id, question)
                self.__doc.add_paragraph()
                id += 1
        
    def save(self) -> None:
        self.__doc.save(self.__file)

class XiaobaosoutiWriter(object):
    
    def __init__(self, file, questions: List[Question]) -> None:
        self.__file = file
        self.__questions = questions
        self.__xiaobao_questions = []
        
    def __write_single(self, question: Question) -> str:
        xiaobao_options = question.options
        xiaobao_answer = ''
        match question.type:
            case QuestionType.single_choice:
                xiaobao_answer = question.answer[0]
            case QuestionType.multiple_choice:
                xiaobao_answer = "".join(question.answer)
            case QuestionType.binary_choice:
                xiaobao_options = ["正确", "错误"]
                xiaobao_answer = "A" if question.answer else "B"
            case QuestionType.short_answer:
                xiaobao_options = [question.answer]
                xiaobao_answer = "A"
        
        return json.dumps({
                    "q": question.question_stem,
                    "a": xiaobao_options,
                    "ans": xiaobao_answer
                }, ensure_ascii=False)
        
    def write(self) -> None:
        for question in self.__questions:
            self.__xiaobao_questions.append(self.__write_single(question))
            
    
    def save(self):
        with open(self.__file,"w+") as f:
            for xiaobao_question in self.__xiaobao_questions:
                f.write(xiaobao_question+"\n")

class KaoshibaoWriter(object):
    def __init__(self, file, questions: List[Question], tpl_file: str="kaoshibao_tpl.xlsx") -> None:
        self.__file = file
        self.__questions = questions
        self.__wb = openpyxl.load_workbook(filename=tpl_file)
        
    def __write_single(self, ws, row, question: Question) -> None:
        
        ws.cell(row = row, column=1, value=question.question_stem)
        
        kaoshibao_type = ''
        kaoshibao_options = question.options
        kaoshibao_answer = ''
        match question.type:
            case QuestionType.single_choice:
                kaoshibao_type = '单选题'
                kaoshibao_answer = question.answer[0]
            case QuestionType.multiple_choice:
                kaoshibao_type = '多选题'
                kaoshibao_answer = "".join(question.answer)
            case QuestionType.binary_choice:
                kaoshibao_type = '判断题'
                kaoshibao_options = ["正确", "错误"]
                kaoshibao_answer = "A" if question.answer else "B"
            case QuestionType.short_answer:
                kaoshibao_type = '简答题'
                kaoshibao_answer = question.answer
        ws.cell(row = row, column=2, value=kaoshibao_type)
        
        for i, opt in enumerate(kaoshibao_options):
            ws.cell(row = row, column=i+3, value=opt)
        ws.cell(row = row, column=11, value=kaoshibao_answer)
        
    def write(self) -> None:
        ws = self.__wb.worksheets[0]
        for i, question in enumerate(self.__questions):
            self.__write_single(ws, i+3, question)
        
    def save(self) -> None:
        self.__wb.save(self.__file)
